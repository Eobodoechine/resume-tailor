"""
FDE DOCX Renderer
=================
Loads the actual FDE template .docx file (pixel-perfect branded resume format),
then surgically replaces content by cloning styled paragraph XML nodes via
lxml deepcopy. Converts the final DOCX to PDF with LibreOffice headless.

Why cloning instead of building from scratch
--------------------------------------------
• Preserves 100% of original styling: fonts, colors, spacing, indents, bullet
  symbols — no approximation errors.
• Flexible: clone bullet nodes as many times as needed.  Template has 4 bullets
  per role; if data has 2 or 6 we clone the prototype exactly that many times.
• Zero maintenance: visual changes are made in Word, not in Python.

Template structure (backend/templates/fde_template.docx)
---------------------------------------------------------
Table 0 — Header (2 rows)
  Row 0 — 3 cells: [teal strip 240] [name+tagline 7560] [contact 4440]
  Row 1 — 1 merged cell: full-width teal separator bar

Table 1 — Body (1 row, 2 cells)
  Cell 0 — 9216 wide — main content (LEFT column)
    child 0: tcPr
    child 1: nested tbl — PROFILE section header bar
    child 2: para — summary/profile text
    child 3: nested tbl — FEATURED PROJECT section header bar
    child 4: para — project title (run 0: bold name, run 1: gray description)
    child 5: para — project URL (italic teal)
    child 6-9: paras — project bullet points
    child 10: nested tbl — PROFESSIONAL EXPERIENCE section header bar
    child 11+: paras — role header / dates / bullets  (repeated per role)
  Cell 1 — 3024 wide — sidebar (RIGHT column, teal bg)
    paras: spacer, CORE SKILLS, categories, bullets, CERTIFICATIONS,
           cert names/details, TRAINING, ..., EDUCATION, degree, detail
"""
from __future__ import annotations

import io
import os
import subprocess
import tempfile
from copy import deepcopy
from typing import Optional, Any

from docx import Document

from renderers.base import ResumeData, Renderer

# ── Paths ──────────────────────────────────────────────────────────────────
_HERE = os.path.dirname(__file__)
TEMPLATE_PATH = os.path.normpath(
    os.path.join(_HERE, '..', 'templates', 'fde_template.docx')
)

# ── XML namespace ──────────────────────────────────────────────────────────
W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'


def _tag(local: str) -> str:
    return f'{{{W}}}{local}'


# ── Low-level XML helpers ──────────────────────────────────────────────────

def _clone(el: Any) -> Any:
    """Deep-copy an lxml element, preserving all attributes and children."""
    return deepcopy(el)


def _set_text(para_el: Any, text: str) -> None:
    """
    Set text in a single-run paragraph.
    Finds the first w:t, sets its text, then blanks out any remaining w:t
    elements in the same paragraph (guards against unexpected extra runs).
    """
    all_t = para_el.findall(f'.//{_tag("t")}')
    if not all_t:
        return
    all_t[0].text = text
    all_t[0].set(XML_SPACE, 'preserve')
    for extra_t in all_t[1:]:
        extra_t.text = ''


def _set_dual_text(para_el: Any, bold_text: str,
                   light_text: Optional[str] = None) -> None:
    """
    Set text in a 2-run paragraph (bold title + gray subtitle).
    Run 0 gets bold_text; run 1 gets ' — {light_text}'.
    If light_text is falsy, run 1 is removed.

    Uses direct-child iteration instead of .//{w:r} to avoid accidentally
    targeting runs nested inside w:hyperlink or other inline elements.
    """
    runs = [child for child in para_el if child.tag == _tag('r')]
    if not runs:
        return

    t0 = runs[0].find(_tag('t'))
    if t0 is not None:
        t0.text = bold_text
        t0.set(XML_SPACE, 'preserve')

    if light_text and len(runs) >= 2:
        t1 = runs[1].find(_tag('t'))
        if t1 is not None:
            t1.text = f' — {light_text}'
            t1.set(XML_SPACE, 'preserve')
    elif not light_text and len(runs) >= 2:
        run1 = runs[1]
        parent = run1.getparent()
        if parent is not None:
            parent.remove(run1)


def _set_bullet_text(para_el: Any, text: str) -> None:
    """
    Set content in a 2-run bullet paragraph.
    Run 0 = '• ' symbol (left untouched). Run 1 = content text.
    Falls back to run 0 if there is only one run.
    """
    runs = para_el.findall(f'.//{_tag("r")}')
    target = runs[1] if len(runs) >= 2 else (runs[0] if runs else None)
    if target is not None:
        t = target.find(_tag('t'))
        if t is not None:
            t.text = text
            t.set(XML_SPACE, 'preserve')


# ── Cell-structure helpers ─────────────────────────────────────────────────

def _section_anchors(tc_el: Any) -> dict:
    """
    Walk direct children of a <w:tc> element and return a dict mapping
    section name → the nested <w:tbl> that holds that section's header bar.
    Keys: 'profile', 'featured', 'experience'
    """
    anchors: dict = {}
    for child in tc_el:
        if child.tag != _tag('tbl'):
            continue
        texts = [t.text for t in child.findall(f'.//{_tag("t")}') if t.text]
        combined = ' '.join(texts)
        if 'PROFILE' in combined:
            anchors['profile'] = child
        elif 'FEATURED PROJECT' in combined:
            anchors['featured'] = child
        elif 'PROFESSIONAL EXPERIENCE' in combined:
            anchors['experience'] = child
    return anchors


def _children_between(tc_el: Any, start_el: Any,
                       end_el: Optional[Any]) -> list:
    """
    Return the direct children of tc_el that appear strictly AFTER start_el
    and strictly BEFORE end_el.  If end_el is None, return everything after
    start_el.
    """
    result: list = []
    inside = False
    for child in tc_el:
        if child is start_el:
            inside = True
            continue
        if end_el is not None and child is end_el:
            break
        if inside:
            result.append(child)
    return result


def _remove_between(tc_el: Any, start_el: Any,
                    end_el: Optional[Any] = None) -> None:
    """Remove direct children after start_el up to (not including) end_el."""
    for el in _children_between(tc_el, start_el, end_el):
        tc_el.remove(el)


def _insert_after(tc_el: Any, anchor_el: Any, elements: list) -> None:
    """Insert elements immediately after anchor_el in tc_el."""
    children = list(tc_el)
    try:
        idx = children.index(anchor_el) + 1
    except ValueError:
        # anchor not found — append to end rather than crash
        idx = len(children)
    for i, el in enumerate(elements):
        tc_el.insert(idx + i, el)


# ── Header ─────────────────────────────────────────────────────────────────

def _fill_header(header_tbl: Any, data: ResumeData) -> None:
    """
    Update name, tagline, and contact items in the 3-column header table.
    Row 0 cells: [teal strip] [name + tagline] [contact list, right-aligned]
    """
    cells = header_tbl.rows[0].cells
    name_cell    = cells[1]   # 7560-wide centre column
    contact_cell = cells[2]   # 4440-wide right column

    # ── Name + tagline ────────────────────────────────────────────────────
    name_paras = name_cell._tc.findall(_tag('p'))
    if name_paras:
        _set_text(name_paras[0], (data.get('name') or '').upper())
    if len(name_paras) > 1:
        _set_text(name_paras[1], data.get('tagline') or '')

    # ── Contact lines ─────────────────────────────────────────────────────
    contact_items = [v for v in [
        data.get('location'),
        data.get('email'),
        data.get('phone'),
        data.get('linkedin'),
        data.get('website'),
        data.get('github'),   # shown separately so both appear if present
    ] if v]

    contact_paras = contact_cell._tc.findall(_tag('p'))
    proto = _clone(contact_paras[0]) if contact_paras else None

    for p in contact_paras:
        contact_cell._tc.remove(p)

    for item in contact_items:
        new_p = _clone(proto)
        _set_text(new_p, item)
        contact_cell._tc.append(new_p)

    # Keep at least one paragraph so the cell is not empty
    if not contact_items and proto is not None:
        contact_cell._tc.append(proto)


# ── Main column ─────────────────────────────────────────────────────────────

def _rebuild_main(main_cell: Any, data: ResumeData) -> None:
    """
    Rebuild the left content column:
      • Summary (PROFILE section)
      • Featured project (optional)
      • Professional experience (N roles, each with M bullets)
    All content is built by cloning styled prototype paragraphs from the
    template, so no formatting is re-coded in Python.
    """
    tc = main_cell._tc
    anchors = _section_anchors(tc)
    profile_tbl  = anchors.get('profile')
    featured_tbl = anchors.get('featured')
    exp_tbl      = anchors.get('experience')

    # ── 1. Summary ────────────────────────────────────────────────────────
    next_after_profile = featured_tbl if featured_tbl is not None else exp_tbl
    if profile_tbl is not None:
        summary_paras = _children_between(tc, profile_tbl, next_after_profile)
        proto_summary = _clone(summary_paras[0]) if summary_paras else None
        _remove_between(tc, profile_tbl, next_after_profile)
        if proto_summary is not None:
            summary_text = data.get('summary') or ''
            # Split on newlines so Claude's intentional line-breaks produce
            # separate styled paragraphs instead of collapsing inside one <w:t>.
            # Empty lines are skipped; fall back to one empty para so the
            # PROFILE section is never left structurally empty.
            lines = [ln.strip() for ln in summary_text.splitlines() if ln.strip()]
            if not lines:
                # summary_text was empty or all-whitespace — write one empty para
                # rather than echoing whitespace into the <w:t> element.
                lines = [""]
            new_paras = []
            for line in lines:
                p = _clone(proto_summary)
                _set_text(p, line)
                new_paras.append(p)
            _insert_after(tc, profile_tbl, new_paras)

    # ── 2. Featured Project ───────────────────────────────────────────────
    fp = data.get('featured_project')

    if featured_tbl is not None:
        fp_paras = _children_between(tc, featured_tbl, exp_tbl)

        # Save prototypes before we wipe the section
        proto_fp_title  = _clone(fp_paras[0]) if len(fp_paras) > 0 else None
        proto_fp_url    = _clone(fp_paras[1]) if len(fp_paras) > 1 else None
        proto_fp_bullet = _clone(fp_paras[2]) if len(fp_paras) > 2 else None

        _remove_between(tc, featured_tbl, exp_tbl)

        if not fp:
            # No featured project data — remove the section header too
            tc.remove(featured_tbl)
        else:
            new_els: list = []

            if proto_fp_title is not None:
                _set_dual_text(proto_fp_title,
                               fp.get('name') or '',
                               fp.get('description') or None)
                new_els.append(proto_fp_title)

            if proto_fp_url is not None and fp.get('url'):
                _set_text(proto_fp_url, fp['url'])
                new_els.append(proto_fp_url)

            if proto_fp_bullet is not None:
                for bullet in (fp.get('bullets') or []):
                    new_b = _clone(proto_fp_bullet)
                    _set_bullet_text(new_b, bullet)
                    new_els.append(new_b)

            _insert_after(tc, featured_tbl, new_els)

    # ── 3. Professional Experience ────────────────────────────────────────
    if exp_tbl is not None:
        exp_paras = _children_between(tc, exp_tbl, None)

        # Prototypes: first role block supplies the three paragraph styles
        proto_role_hdr = _clone(exp_paras[0]) if len(exp_paras) > 0 else None
        proto_date     = _clone(exp_paras[1]) if len(exp_paras) > 1 else None
        proto_bullet   = _clone(exp_paras[2]) if len(exp_paras) > 2 else None

        _remove_between(tc, exp_tbl, None)

        new_els = []
        for role in (data.get('experience') or []):
            # Role header: "Title — Company — Location"
            if proto_role_hdr is not None:
                new_hdr = _clone(proto_role_hdr)
                company_loc = ' — '.join(
                    filter(None, [role.get('company'), role.get('location')])
                )
                _set_dual_text(new_hdr,
                               role.get('title') or '',
                               company_loc or None)
                new_els.append(new_hdr)

            # Date line (italic teal)
            if proto_date is not None:
                new_date = _clone(proto_date)
                _set_text(new_date, role.get('dates') or '')
                new_els.append(new_date)

            # Bullets — clone one node per bullet item
            if proto_bullet is not None:
                for bullet in (role.get('bullets') or []):
                    new_b = _clone(proto_bullet)
                    _set_bullet_text(new_b, bullet)
                    new_els.append(new_b)

        _insert_after(tc, exp_tbl, new_els)


# ── Sidebar ─────────────────────────────────────────────────────────────────

def _find_sidebar_protos(all_paras: list) -> dict:
    """
    Extract prototype paragraph elements from the sidebar using a mix of
    positional lookup (for the fixed skills block at the top) and text-scan
    (for CERTIFICATIONS and EDUCATION which follow variable-length skill lists).
    Returns a dict of deepcopy'd lxml elements keyed by role name.

    Actual sidebar paragraph order (index 0 = first <w:p> child of sidebar tc):
      [0]  "CORE SKILLS"          → skill_section
      [1]  "AI & ML"              → category
      [2]  "• Vertex AI..."       → skill_bullet  (2 runs: bullet symbol + content)
      ...  (more bullets + more groups)
      [N]  "CERTIFICATIONS"       → cert_section  (found by text scan)
      [N+1] cert name             → cert_name
      [N+2] cert detail           → cert_detail
      ...
      [M]  "EDUCATION"            → edu_section   (found by text scan)
      [M+1] degree line           → edu_degree    (2 runs: degree + school)
      [M+2] detail line           → edu_detail
    """
    protos: dict = {}

    # Fixed positions at top of sidebar (no leading spacer in this template)
    if len(all_paras) > 0:
        protos['skill_section'] = _clone(all_paras[0])   # "CORE SKILLS"
    if len(all_paras) > 1:
        protos['category']      = _clone(all_paras[1])   # "AI & ML"
    if len(all_paras) > 2:
        protos['skill_bullet']  = _clone(all_paras[2])   # "• Vertex AI..." (2 runs)

    # Scan for CERTIFICATIONS and EDUCATION by their header text
    cert_idx: Optional[int] = None
    edu_idx:  Optional[int] = None

    for i, p in enumerate(all_paras):
        texts = [t.text for t in p.findall(f'.//{_tag("t")}') if t.text]
        text0 = texts[0] if texts else ''
        if 'CERTIFICATIONS' in text0 and cert_idx is None:
            cert_idx = i
        if 'EDUCATION' in text0 and edu_idx is None:
            edu_idx = i

    if cert_idx is not None:
        protos['cert_section'] = _clone(all_paras[cert_idx])
        if cert_idx + 1 < len(all_paras):
            protos['cert_name']   = _clone(all_paras[cert_idx + 1])
        if cert_idx + 2 < len(all_paras):
            protos['cert_detail'] = _clone(all_paras[cert_idx + 2])

    if edu_idx is not None:
        protos['edu_section'] = _clone(all_paras[edu_idx])
        if edu_idx + 1 < len(all_paras):
            protos['edu_degree']  = _clone(all_paras[edu_idx + 1])
        if edu_idx + 2 < len(all_paras):
            protos['edu_detail']  = _clone(all_paras[edu_idx + 2])

    return protos


def _rebuild_sidebar(sidebar_cell: Any, data: ResumeData) -> None:
    """
    Rebuild the right teal sidebar:
      Core Skills → Certifications → Education
    All paragraph styling comes from cloned template prototypes.
    """
    tc = sidebar_cell._tc
    all_paras = tc.findall(_tag('p'))
    protos = _find_sidebar_protos(all_paras)

    # Wipe all existing paragraphs
    for p in all_paras:
        tc.remove(p)

    def add(key: str,
            text: Optional[str] = None,
            bold_text: Optional[str] = None,
            light_text: Optional[str] = None,
            bullet_text: Optional[str] = None) -> None:
        proto = protos.get(key)
        if proto is None:
            return
        el = _clone(proto)
        if text is not None:
            _set_text(el, text)
        elif bold_text is not None:
            _set_dual_text(el, bold_text, light_text)
        elif bullet_text is not None:
            _set_bullet_text(el, bullet_text)
        tc.append(el)

    # Core Skills
    add('skill_section')
    for group in (data.get('skills') or []):
        add('category', text=group.get('category') or '')
        for item in (group.get('items') or []):
            add('skill_bullet', bullet_text=item)

    # Certifications (includes training entries if Claude lumps them together)
    certs = data.get('certifications') or []
    if certs and protos.get('cert_section') is not None:
        add('cert_section')
        for cert in certs:
            if isinstance(cert, dict):
                name   = cert.get('name') or str(cert)
                detail = cert.get('detail')
            else:
                name   = str(cert)
                detail = None
            add('cert_name', text=name)
            if detail and protos.get('cert_detail') is not None:
                add('cert_detail', text=detail)

    # Education
    education = data.get('education') or []
    if education and protos.get('edu_section') is not None:
        add('edu_section')
        for edu in education:
            if isinstance(edu, dict):
                degree = edu.get('degree') or ''
                school = edu.get('school')
                detail = edu.get('detail')
            else:
                degree = str(edu)
                school = None
                detail = None
            if protos.get('edu_degree') is not None:
                add('edu_degree', bold_text=degree, light_text=school)
            if detail and protos.get('edu_detail') is not None:
                add('edu_detail', text=detail)


# ── LibreOffice PDF conversion ─────────────────────────────────────────────

def _docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes → PDF bytes using LibreOffice headless."""
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, 'resume.docx')
        pdf_path  = os.path.join(tmp, 'resume.pdf')

        with open(docx_path, 'wb') as fh:
            fh.write(docx_bytes)

        env = {**os.environ, 'HOME': '/tmp'}
        result = subprocess.run(
            ['libreoffice', '--headless', '--convert-to', 'pdf',
             '--outdir', tmp, docx_path],
            capture_output=True,
            timeout=120,   # 60s was too tight for first cold-start on Render free tier
            env=env,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f'LibreOffice failed (exit {result.returncode}): '
                f'{result.stderr.decode(errors="replace")}'
            )
        if not os.path.exists(pdf_path):
            raise RuntimeError(
                f'LibreOffice exited 0 but no PDF produced. '
                f'stdout: {result.stdout.decode(errors="replace")}'
            )
        with open(pdf_path, 'rb') as fh:
            return fh.read()


# ── Public Renderer ────────────────────────────────────────────────────────

class FDEDocxRenderer:
    """
    Pixel-perfect FDE resume renderer using lxml deepcopy node cloning.

    Reads the template bytes once at class load time and re-parses them into
    a fresh Document on every render() call.  Re-parsing from in-memory bytes
    is much faster than reading from disk and guarantees complete state isolation
    between concurrent requests — python-docx Document objects cannot be safely
    deepcopy'd across calls because the underlying lxml element tree may share
    references.

    The template supplies all visual styling; this class only handles data
    placement and node count.
    """

    # Read template bytes once — re-parse per render() for full isolation.
    _template_bytes: bytes = open(TEMPLATE_PATH, "rb").read()

    def render(self, data: ResumeData) -> bytes:
        doc = Document(io.BytesIO(self._template_bytes))

        header_tbl   = doc.tables[0]   # 2-row header table
        body_tbl     = doc.tables[1]   # 1-row body table
        main_cell    = body_tbl.rows[0].cells[0]   # left — main content
        sidebar_cell = body_tbl.rows[0].cells[1]   # right — teal sidebar

        _fill_header(header_tbl, data)
        _rebuild_main(main_cell, data)
        _rebuild_sidebar(sidebar_cell, data)

        buf = io.BytesIO()
        doc.save(buf)
        return _docx_to_pdf(buf.getvalue())
