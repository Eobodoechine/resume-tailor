"""
Extract plain text from PDF and DOCX resume files.
"""
import io
import logging
import pdfplumber
import docx

logger = logging.getLogger(__name__)


def extract_text(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        return _extract_pdf(file_bytes, filename_hint=filename)
    elif ext == "docx":
        return _extract_docx(file_bytes, filename_hint=filename)
    elif ext == "doc":
        # python-docx only reads OOXML (.docx); legacy OLE2 .doc is unsupported (B4).
        raise ValueError("Legacy .doc files are not supported — please save as .docx or PDF.")
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _extract_pdf(file_bytes: bytes, filename_hint: str = "") -> str:
    text_parts = []
    pages_total = 0
    pages_empty = 0
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            pages_total += 1
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            else:
                pages_empty += 1
    text = "\n\n".join(text_parts)
    if not text.strip():
        logger.warning(
            "[extractor] _extract_pdf zero chars extracted — possible image-only or scanned PDF  "
            "filename=%r  pages_total=%d  pages_empty=%d",
            filename_hint, pages_total, pages_empty,
        )
    else:
        logger.info(
            "[extractor] _extract_pdf OK  filename=%r  chars=%d  pages_total=%d  pages_empty=%d",
            filename_hint, len(text), pages_total, pages_empty,
        )
    return text


def _extract_docx(file_bytes: bytes, filename_hint: str = "") -> str:
    """
    Extract text from a DOCX file, including paragraphs AND tables.

    A huge fraction of real-world resumes are built in two-column Word
    tables — extracting only doc.paragraphs would silently drop that
    content. We also pull text from section headers/footers in case the
    candidate put their contact info there.

    Merged table cells in python-docx repeat the same underlying _tc XML
    element across multiple row.cells entries. We track visited _tc ids to
    deduplicate merged cells and avoid repeating their text (TD-08).
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    parts: list[str] = []
    paragraph_count = 0

    for p in doc.paragraphs:
        paragraph_count += 1
        text = p.text.strip()
        if text:
            parts.append(text)

    seen_cell_tcs: set[int] = set()
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id in seen_cell_tcs:
                    continue
                seen_cell_tcs.add(tc_id)
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    for section in doc.sections:
        for source in (section.header, section.footer):
            for p in source.paragraphs:
                text = p.text.strip()
                if text:
                    parts.append(text)

    text = "\n".join(parts)
    if not text.strip():
        logger.warning(
            "[extractor] _extract_docx zero chars extracted — possible empty or image-only DOCX  "
            "filename=%r  paragraphs=%d",
            filename_hint, paragraph_count,
        )
    else:
        logger.info(
            "[extractor] _extract_docx OK  filename=%r  chars=%d  paragraphs=%d",
            filename_hint, len(text), paragraph_count,
        )
    return text
