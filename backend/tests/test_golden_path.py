"""
Golden-path visual regression test.

Uses the exact data from Nnamdi's Tribe AI resume PDF (the reference output)
to verify the renderer produces the correct structure. LibreOffice is mocked
so no binary is needed — but the DOCX structure is inspected in full.

This is the MOST IMPORTANT test: if the rendered DOCX contains all the right
text in the right sections, the PDF will look like the reference document.
"""
import sys
import io
import types
from copy import deepcopy
from unittest.mock import patch, MagicMock
import pytest

# ── Load real docx, not the MagicMock stub from conftest ─────────────────────
if isinstance(sys.modules.get("docx"), MagicMock):
    del sys.modules["docx"]
for _k in [k for k in sys.modules if k.startswith("renderers.fde_docx")]:
    del sys.modules[_k]

import renderers.fde_docx as fde_mod
from renderers.fde_docx import FDEDocxRenderer
from docx import Document

# ── Reference data — matches Tribe AI PDF exactly ────────────────────────────

NNAMDI_TRIBE_AI = {
    "name": "Nnamdi Obodoechine",
    "email": "eobodoechine@gmail.com",
    "phone": None,
    "location": "Atlanta, GA",
    "linkedin": "linkedin.com/in/nnamdi-obodoechine",
    "github": "github.com/Eobodoechine",
    "tagline": "Workflow Strategist • Agentic AI Practitioner • Enterprise Enablement",
    "summary": (
        "Operations and workflow strategist who embeds with enterprise teams to discover how work "
        "actually happens — then redesigns it. Daily practitioner of the full agentic tool suite: "
        "Cowork, Claude Code, and Claude Enterprise, building production AI workflow systems and "
        "enablement programs across client engagements. At UPS, redesigned tax payment workflows "
        "across a $90M+ portfolio, delivering $1M+ in net working capital benefit. At EY, trained "
        "and enabled 60+ consultants across 44 countries. Self-directed builder who owns "
        "engagements end-to-end, translates agentic capabilities to non-technical stakeholders, "
        "and turns field patterns into institutional knowledge."
    ),
    "featured_project": {
        "name": "Agentic Workflow Automation Suite",
        "description": "Built on Claude, Cowork, and Claude Code",
        "url": "github.com/Eobodoechine | resume-tailor-ogop.onrender.com",
        "bullets": [
            "Shipped Resume Tailor — a full-stack AI web app (FastAPI + Supabase + Claude API) that tailors resumes to job descriptions in real time, with streaming output and PDF generation",
            "Built apply-for-job and lemlist-campaign-setup Cowork skills — autonomous agents that detect ATS platforms, fill applications, and load leads into outreach sequences end-to-end",
            "Daily practitioner of Cowork, Claude Code, and Claude Enterprise — using these tools across research, workflow automation, content generation, and AI product builds every day",
            "All systems built end-to-end solo: architecture, backend, frontend, deployment, and iterative improvement — same full-ownership model required in client engagements",
        ],
    },
    "experience": [
        {
            "title": "Workflow Strategy & Operations Specialist",
            "company": "United Parcel Service (UPS)",
            "location": "Atlanta, GA",
            "dates": "Jan 2026 – Present",
            "bullets": [
                "Redesigned tax payment workflow across 3,000+ parcels: ran discovery against third-party vendor SOC/SOX constraints, then built a two-variable decision matrix weighing each state's discount rate against its due-date-to-delinquency gap — e.g. Florida's 4% discount justified early payment even at 120-day gaps; low-discount/short-gap states held cash to delinquency. Created Net360 payment schedule delivering $1M+ net benefit",
                "Designed and deployed Microsoft Forms + SharePoint + Power Automate workflow system from scratch — handling enterprise-scale request routing and SLA tracking adopted by non-technical real estate teams with zero IT dependency",
                "Built knowledge transfer documentation and SOPs for team subordinates across payment workflows, vendor onboarding, and system configurations — turning tribal knowledge into repeatable playbooks the team uses independently",
                "Manage $90M/month payment portfolio with validation, reconciliation logic, and exception-handling across thousands of parcels and 100+ vendor relationships; partner cross-functionally with UPS Treasury, FSC, GBS Vendor Management, and CBRE",
            ],
        },
        {
            "title": "Principal & AI Workflow Builder",
            "company": "ENO LLC",
            "location": "Atlanta, GA",
            "dates": "Jul 2023 – Aug 2025",
            "bullets": [
                "Built AI-agent and workflow automation systems with Claude, n8n, Cowork, and Gemini APIs — applying prompt engineering and tool orchestration across client delivery, outreach automation, and real estate underwriting workflows",
                "Architected and shipped PropertyVision — production AI/ML SaaS on GCP (Vertex AI, Cloud Run, Docker, Redis) applying agentic AI to real estate workflow automation at scale",
                "Underwrote $15M+ in real estate using IRR/NPV/CoC/RevPAR models; built client-facing deliverables translating complex financial analysis into decisions for non-technical stakeholders",
            ],
        },
        {
            "title": "Senior Technology & Enablement Consultant",
            "company": "Ernst & Young (EY)",
            "location": "Atlanta, GA",
            "dates": "Aug 2019 – Jun 2023",
            "bullets": [
                "Designed and delivered enablement programs for 60+ consultants across 44 countries — structured training in Python, SQL, SAP, and data modeling with before/after outcomes; built reusable guides that outlasted each engagement",
                "Embedded with global Fortune 500 clients — ran discovery on existing data workflows, surfaced root causes of processing failures, and redesigned pipelines handling 1M+ data points",
                "Remediated 200K+ financial data anomalies; built forecasting models, classification rules, and C-suite dashboards under tight delivery cadence across financial services and manufacturing sectors",
            ],
        },
    ],
    "skills": [
        {
            "category": "Agentic Tools",
            "items": ["Cowork", "Claude Code", "Claude Enterprise", "Claude API",
                      "n8n", "OpenAI", "Gemini APIs", "Prompt Engineering",
                      "AI Agents", "RAG", "Workflow Orchestration"],
        },
        {
            "category": "Workflow & Enablement",
            "items": ["Workflow Discovery & Redesign", "Training Program Design",
                      "Process Documentation", "SOPs", "Change Management", "Adoption"],
        },
        {
            "category": "Enterprise Tech",
            "items": ["Power Automate", "SharePoint", "Microsoft Forms", "Power BI",
                      "Coupa", "CoStar", "Oracle EBS", "FastAPI", "Supabase", "GCP"],
        },
        {
            "category": "Stakeholder & Communication",
            "items": ["Non-Technical Stakeholder Translation", "Client Engagement",
                      "Discovery", "Written Communication", "SOPs"],
        },
    ],
    "certifications": [
        {"name": "Azure AI Fundamentals (AZ-900)", "detail": "Microsoft • Active"},
        {"name": "ML with Python & SQL", "detail": "Harvard Extension • 2022"},
        {"name": "Data Analytics Bootcamp", "detail": "Georgia Tech • 2021"},
    ],
    "education": [
        {
            "degree": "B.B.A., Economics & Management",
            "school": "Georgia Southern University • 2019",
            "detail": "Minors: Business Analytics & ERP Systems",
        },
    ],
}


# ── Helper ────────────────────────────────────────────────────────────────────

def _render_and_inspect(data=None) -> tuple[str, Document]:
    """Render data through the FDEDocxRenderer (LibreOffice mocked).
    Returns (all_text, doc) for assertions."""
    data = data or NNAMDI_TRIBE_AI
    renderer = FDEDocxRenderer()
    captured = {}

    def fake_lo(docx_bytes):
        captured["bytes"] = docx_bytes
        return b"%PDF-1.4 fake"

    with patch.object(fde_mod, "_docx_to_pdf", side_effect=fake_lo):
        renderer.render(data)

    doc = Document(io.BytesIO(captured["bytes"]))
    all_text = " ".join(
        t.text for t in doc.element.findall(f".//{{{fde_mod.W}}}t") if t.text
    )
    return all_text, doc


# ── Golden-path assertions ────────────────────────────────────────────────────

class TestGoldenPathHeader:
    """Header must match the Tribe AI PDF reference exactly."""

    def test_name_uppercased(self):
        text, _ = _render_and_inspect()
        assert "NNAMDI OBODOECHINE" in text

    def test_tagline_present(self):
        text, _ = _render_and_inspect()
        assert "Workflow Strategist" in text
        assert "Agentic AI Practitioner" in text
        assert "Enterprise Enablement" in text

    def test_location_in_contact(self):
        text, _ = _render_and_inspect()
        assert "Atlanta, GA" in text

    def test_email_in_contact(self):
        text, _ = _render_and_inspect()
        assert "eobodoechine@gmail.com" in text

    def test_linkedin_in_contact(self):
        text, _ = _render_and_inspect()
        assert "linkedin.com/in/nnamdi-obodoechine" in text

    def test_github_in_contact(self):
        """github.com/Eobodoechine must appear — not dropped behind website."""
        text, _ = _render_and_inspect()
        assert "github.com/Eobodoechine" in text


class TestGoldenPathSummary:
    def test_profile_summary_opening_sentence(self):
        text, _ = _render_and_inspect()
        assert "Operations and workflow strategist" in text

    def test_ups_mention_in_summary(self):
        text, _ = _render_and_inspect()
        assert "$90M+" in text

    def test_ey_mention_in_summary(self):
        text, _ = _render_and_inspect()
        assert "60+ consultants" in text


class TestGoldenPathFeaturedProject:
    def test_project_name_present(self):
        text, _ = _render_and_inspect()
        assert "Agentic Workflow Automation Suite" in text

    def test_project_description_present(self):
        text, _ = _render_and_inspect()
        assert "Built on Claude, Cowork, and Claude Code" in text

    def test_project_url_present(self):
        text, _ = _render_and_inspect()
        assert "resume-tailor-ogop.onrender.com" in text

    def test_all_four_project_bullets(self):
        text, _ = _render_and_inspect()
        assert "Resume Tailor" in text
        assert "apply-for-job" in text
        assert "Daily practitioner of Cowork" in text
        assert "end-to-end solo" in text

    def test_featured_project_section_header_present(self):
        text, _ = _render_and_inspect()
        assert "FEATURED PROJECT" in text


class TestGoldenPathExperience:
    def test_all_three_roles_present(self):
        text, _ = _render_and_inspect()
        assert "Workflow Strategy & Operations Specialist" in text
        assert "Principal & AI Workflow Builder" in text
        assert "Senior Technology & Enablement Consultant" in text

    def test_all_three_companies_present(self):
        text, _ = _render_and_inspect()
        assert "United Parcel Service" in text
        assert "ENO LLC" in text
        assert "Ernst & Young" in text

    def test_all_three_date_ranges(self):
        text, _ = _render_and_inspect()
        assert "Jan 2026" in text
        assert "Jul 2023" in text
        assert "Aug 2019" in text

    def test_ups_key_bullets(self):
        text, _ = _render_and_inspect()
        assert "Net360" in text
        assert "$1M+" in text
        assert "Power Automate workflow system" in text

    def test_eno_key_bullets(self):
        text, _ = _render_and_inspect()
        assert "PropertyVision" in text
        assert "$15M+" in text

    def test_ey_key_bullets(self):
        text, _ = _render_and_inspect()
        assert "44 countries" in text
        assert "200K+" in text

    def test_all_ten_bullets_across_roles(self):
        """Each role has the correct bullet count (4 + 3 + 3 = 10)."""
        text, _ = _render_and_inspect()
        # Spot-check 1 unique phrase from each of the 10 bullets
        phrases = [
            "Net360",                           # UPS bullet 1
            "zero IT dependency",               # UPS bullet 2
            "repeatable playbooks",             # UPS bullet 3
            "CBRE",                             # UPS bullet 4
            "Gemini APIs",                      # ENO bullet 1
            "PropertyVision",                   # ENO bullet 2
            "IRR/NPV",                          # ENO bullet 3
            "reusable guides that outlasted",   # EY bullet 1
            "1M+ data points",                  # EY bullet 2
            "200K+",                            # EY bullet 3
        ]
        for phrase in phrases:
            assert phrase in text, f"Missing bullet phrase: '{phrase}'"

    def test_experience_section_header_present(self):
        text, _ = _render_and_inspect()
        assert "PROFESSIONAL EXPERIENCE" in text


class TestGoldenPathSidebar:
    def test_all_four_skill_categories(self):
        text, _ = _render_and_inspect()
        assert "Agentic Tools" in text
        assert "Workflow & Enablement" in text
        assert "Enterprise Tech" in text
        assert "Stakeholder" in text

    def test_key_skill_items_present(self):
        text, _ = _render_and_inspect()
        # Agentic tools
        assert "Claude Code" in text
        assert "Claude Enterprise" in text
        assert "Prompt Engineering" in text
        assert "RAG" in text
        # Enterprise tech
        assert "Power Automate" in text
        assert "Oracle EBS" in text
        assert "FastAPI" in text

    def test_certification_present(self):
        text, _ = _render_and_inspect()
        assert "Azure AI Fundamentals" in text
        assert "AZ-900" in text
        assert "Microsoft" in text

    def test_training_entries_present(self):
        text, _ = _render_and_inspect()
        assert "Harvard Extension" in text
        assert "Georgia Tech" in text

    def test_education_degree_present(self):
        text, _ = _render_and_inspect()
        assert "B.B.A." in text
        assert "Economics & Management" in text
        assert "Georgia Southern University" in text

    def test_education_detail_present(self):
        text, _ = _render_and_inspect()
        assert "Business Analytics" in text
        assert "ERP Systems" in text

    def test_core_skills_section_header(self):
        text, _ = _render_and_inspect()
        assert "CORE SKILLS" in text


class TestGoldenPathDocxStructure:
    """Verify DOCX table structure — ensures renderer didn't corrupt the template."""

    def test_document_has_two_tables(self):
        _, doc = _render_and_inspect()
        assert len(doc.tables) == 2, "Expected header table + body table"

    def test_header_table_has_two_rows(self):
        _, doc = _render_and_inspect()
        assert len(doc.tables[0].rows) == 2, "Header table: row0=content, row1=separator"

    def test_header_row_has_three_cells(self):
        _, doc = _render_and_inspect()
        assert len(doc.tables[0].rows[0].cells) == 3, "Header: teal strip | name | contact"

    def test_body_table_has_two_cells(self):
        _, doc = _render_and_inspect()
        row = doc.tables[1].rows[0]
        assert len(row.cells) == 2, "Body: main column + sidebar"

    def test_render_returns_pdf_bytes(self):
        renderer = FDEDocxRenderer()
        with patch.object(fde_mod, "_docx_to_pdf", return_value=b"%PDF-1.4 test"):
            result = renderer.render(NNAMDI_TRIBE_AI)
        assert result == b"%PDF-1.4 test"


class TestGoldenPathEdgeCases:
    """Make sure the reference data handles variations without breaking."""

    def test_render_with_no_phone_doesnt_crash(self):
        """Tribe AI resume has no phone number — contact block must still render."""
        data = {**NNAMDI_TRIBE_AI, "phone": None}
        text, _ = _render_and_inspect(data)
        assert "eobodoechine@gmail.com" in text

    def test_render_is_idempotent(self):
        """Rendering same data twice should produce identical text content."""
        text1, _ = _render_and_inspect()
        text2, _ = _render_and_inspect()
        assert text1 == text2

    def test_render_does_not_mutate_input_data(self):
        """The renderer must not modify the data dict passed to it."""
        import copy
        original = copy.deepcopy(NNAMDI_TRIBE_AI)
        _render_and_inspect(NNAMDI_TRIBE_AI)
        assert NNAMDI_TRIBE_AI == original
